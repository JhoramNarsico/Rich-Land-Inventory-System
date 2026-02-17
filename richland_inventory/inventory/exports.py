import csv
import io
import os
from decimal import Decimal

from django.http import HttpResponse
from django.conf import settings
from django.utils import timezone
from django.utils.text import slugify

# Excel
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.worksheet.page import PageMargins
from openpyxl.drawing.image import Image

# Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .utils import render_to_pdf

# --- HELPERS ---

def setup_word_document_margins(document):
    """Sets very narrow margins (0.25 inch) and Letter size (8.5x11) for the document."""
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)

def set_cell_background(cell, color_hex):
    """Helper to set background color of a Word table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_excel_logo(ws):
    """Adds the company logo to the top-left of the worksheet."""
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'images', 'logo.png')
    if os.path.exists(logo_path):
        img = Image(logo_path)
        img.width = 60
        img.height = 60
        ws.add_image(img, 'A1')

def create_header(document, title, subtitle_lines=[], width_inches=7.8):
    """Creates a professional header with Company Name and Report Details."""
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    # Usable width approx width_inches (leaving slight buffer).
    table.columns[0].width = Inches(width_inches / 2)
    table.columns[1].width = Inches(width_inches / 2)
    
    # Left: Company Name
    cell_left = table.cell(0, 0)
    p = cell_left.paragraphs[0]
    
    # Logo
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'images', 'logo.png')
    if os.path.exists(logo_path):
        run_logo = p.add_run()
        run_logo.add_picture(logo_path, width=Inches(0.6))
        p.add_run("  ")

    run = p.add_run("Rich Land Auto Supply")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Right: Report Info
    cell_right = table.cell(0, 1)
    p = cell_right.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x27, 0x99, 0xA5)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for line in subtitle_lines:
        p2 = cell_right.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = p2.add_run(line)
        run2.font.name = 'Arial'
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    
    document.add_paragraph() # Spacer

def style_table_header(row, headers):
    """Styles the header row of a table."""
    for i, text in enumerate(headers):
        cell = row.cells[i]
        set_cell_background(cell, "2C3E50")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        run.bold = True
        run.font.size = Pt(10)

# --- EXPORT FUNCTIONS ---

def generate_sow_history_export(customer, sows, format_type, request):
    filename = f"SOW_History_{slugify(customer.name)}_{timezone.now().strftime('%Y%m%d')}"

    if format_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        writer = csv.writer(response)
        writer.writerow(['Date', 'Application', 'Hose Type', 'Diameter', 'Length', 'Pressure', 'Fitting A', 'Fitting B', 'Notes'])
        for sow in sows:
            writer.writerow([
                sow.date_created.strftime('%Y-%m-%d'),
                sow.application,
                sow.hose_type,
                sow.diameter,
                sow.length,
                sow.pressure,
                sow.fitting_a,
                sow.fitting_b,
                sow.notes
            ])
        return response

    elif format_type == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.title = "SOW History"
        add_excel_logo(ws)
        ws['B1'] = "SOW HISTORY"
        ws['B1'].font = Font(bold=True, size=12)
        ws['C1'] = f"Customer: {customer.name}"
        ws['D1'] = f"Date: {timezone.now().strftime('%Y-%m-%d')}"
        ws.append([])
        headers = ['Date', 'Application', 'Hose Type', 'Diameter', 'Length', 'Pressure', 'Fitting A', 'Fitting B', 'Notes']
        ws.append(headers)
        for sow in sows:
            ws.append([
                sow.date_created.strftime('%Y-%m-%d'),
                sow.application,
                sow.hose_type,
                sow.diameter,
                sow.length,
                sow.pressure,
                sow.fitting_a,
                sow.fitting_b,
                sow.notes
            ])
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
        wb.save(response)
        return response

    elif format_type == 'word':
        document = Document()
        setup_word_document_margins(document)
        
        create_header(document, "HYDRAULIC SOW HISTORY", [
            f"Customer: {customer.name}",
            f"Generated: {timezone.now().strftime('%B %d, %Y')}"
        ])

        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False 
        # Set approximate widths (Total 7.8 inches)
        table.columns[0].width = Inches(1.0) # Date
        table.columns[1].width = Inches(1.5) # App
        table.columns[2].width = Inches(2.0) # Hose
        table.columns[3].width = Inches(1.6) # Fittings
        table.columns[4].width = Inches(1.7) # Notes

        hdr_cells = table.rows[0].cells
        headers = ['Date', 'Application', 'Hose Details', 'Fittings', 'Notes']
        style_table_header(table.rows[0], headers)

        for sow in sows:
            row_cells = table.add_row().cells
            row_cells[0].text = sow.date_created.strftime('%Y-%m-%d')
            row_cells[1].text = sow.application
            row_cells[2].text = f"{sow.hose_type}\n{sow.diameter}\" x {sow.length}mm\n{sow.pressure} PSI"
            row_cells[3].text = f"A: {sow.fitting_a}\nB: {sow.fitting_b}"
            row_cells[4].text = sow.notes

        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    else: # PDF
        context = {
            'customer': customer,
            'sows': sows,
            'today': timezone.now(),
        }
        pdf = render_to_pdf('inventory/sow_history_pdf.html', context, request=request)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
            return response
        return HttpResponse("Error Generating PDF", status=500)

def generate_expense_report(expenses, format_type, request):
    filename = f"Expense_Report_{timezone.now().strftime('%Y%m%d')}"

    if format_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        writer = csv.writer(response)
        writer.writerow(['Date', 'Category', 'Description', 'Amount', 'Recorded By'])
        for expense in expenses:
            writer.writerow([
                expense.expense_date,
                expense.category.name if expense.category else 'N/A',
                expense.description,
                expense.amount,
                expense.recorded_by.username if expense.recorded_by else 'N/A'
            ])
        return response
    
    elif format_type == 'word':
        document = Document()
        setup_word_document_margins(document)
        
        total = sum(e.amount for e in expenses)
        create_header(document, "EXPENSE REPORT", [
            f"Total Expenses: {total:,.2f}",
            f"Generated: {timezone.now().strftime('%B %d, %Y')}"
        ])

        table = document.add_table(rows=1, cols=4, style='Table Grid')
        table.autofit = False
        table.columns[0].width = Inches(1.1) # Date
        table.columns[1].width = Inches(1.6) # Category
        table.columns[2].width = Inches(3.6) # Description
        table.columns[3].width = Inches(1.5) # Amount

        headers = ['Date', 'Category', 'Description', 'Amount']
        style_table_header(table.rows[0], headers)
        
        for expense in expenses:
            row_cells = table.add_row().cells
            row_cells[0].text = expense.expense_date.strftime('%Y-%m-%d')
            row_cells[1].text = expense.category.name if expense.category else ''
            row_cells[2].text = expense.description
            row_cells[3].text = f"{expense.amount:,.2f}"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    elif format_type == 'pdf':
        context = {
            'expenses': expenses,
            'total_expenses': sum(e.amount for e in expenses),
            'today': timezone.now(),
        }
        pdf = render_to_pdf('inventory/expense_report_pdf.html', context, request=request)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
            return response
    
    return None

def generate_customer_list_export(customers, format_type, request):
    filename = f"Customer_List_{timezone.now().strftime('%Y%m%d')}"

    if format_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        writer = csv.writer(response)
        writer.writerow(['Name', 'Email', 'Phone', 'Address', 'Tax ID', 'Credit Limit', 'Current Balance'])
        for customer in customers:
            writer.writerow([
                customer.name, customer.email, customer.phone, customer.address,
                customer.tax_id, customer.credit_limit, customer.get_balance()
            ])
        return response

    elif format_type == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.title = "Customers"
        add_excel_logo(ws)
        ws['B1'] = 'Customer List'
        ws['B1'].font = Font(bold=True, size=12)
        ws['C1'] = f"Date: {timezone.now().strftime('%Y-%m-%d')}"
        ws.append([])
        headers = ['Name', 'Email', 'Phone', 'Address', 'Tax ID', 'Credit Limit', 'Current Balance']
        ws.append(headers)
        for customer in customers:
            ws.append([
                customer.name, customer.email, customer.phone, customer.address,
                customer.tax_id, f"{customer.credit_limit:.2f}", f"{customer.get_balance():.2f}"
            ])
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
        wb.save(response)
        return response

    elif format_type == 'word':
        document = Document()
        setup_word_document_margins(document)
        
        create_header(document, "CUSTOMER LIST", [
            f"Generated: {timezone.now().strftime('%B %d, %Y')}"
        ])
        
        table = document.add_table(rows=1, cols=4, style='Table Grid')
        table.autofit = False
        table.columns[0].width = Inches(1.8) # Name
        table.columns[1].width = Inches(1.8) # Contact
        table.columns[2].width = Inches(2.6) # Address
        table.columns[3].width = Inches(1.6) # Credit Info

        style_table_header(table.rows[0], ['Name', 'Contact', 'Address', 'Credit Info'])
        
        for customer in customers:
            row_cells = table.add_row().cells
            row_cells[0].text = customer.name
            row_cells[1].text = f"{customer.email or 'N/A'}\n{customer.phone or 'N/A'}"
            row_cells[2].text = customer.address or 'N/A'
            row_cells[3].text = f"Limit: {customer.credit_limit:,.2f}\nBalance: {customer.get_balance():,.2f}"
        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    elif format_type == 'pdf':
        context = {'customers': customers, 'today': timezone.now(), 'request': request}
        pdf = render_to_pdf('inventory/customer_list_pdf.html', context)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
            return response
    
    return None

def generate_customer_statement(customer, final_data, running_balance, format_type, request):
    filename = f"Statement_{slugify(customer.name)}_{timezone.now().strftime('%Y%m%d')}"

    if format_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        writer = csv.writer(response)
        writer.writerow(['Date', 'Reference', 'Description', 'Charge', 'Payment', 'Balance'])
        for row in final_data:
            charge = row['amount'] if not row['is_credit'] else ''
            pay = row['amount'] if row['is_credit'] else ''
            writer.writerow([row['date'].strftime('%Y-%m-%d'), row['ref'], row['desc'], charge, pay, row['balance']])
        return response

    elif format_type == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.title = "Statement"
        add_excel_logo(ws)
        ws['B1'] = "BILLING STATEMENT"; ws['B1'].font = Font(bold=True, size=12)
        ws['B2'] = f"Customer: {customer.name}"
        ws['B3'] = f"Date: {timezone.now().strftime('%Y-%m-%d')}"
        ws.append([]) 
        headers = ['Date', 'Reference', 'Description', 'Charge', 'Payment', 'Balance']
        ws.append(headers)
        for row in final_data:
            charge = row['amount'] if not row['is_credit'] else None
            pay = row['amount'] if row['is_credit'] else None
            ws.append([row['date'].strftime('%Y-%m-%d'), row['ref'], row['desc'], charge, pay, row['balance']])
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
        wb.save(response)
        return response

    elif format_type == 'word':
        document = Document()
        setup_word_document_margins(document)
        
        create_header(document, "BILLING STATEMENT", [
            f"Customer: {customer.name}",
            f"Address: {customer.address or 'N/A'}",
            f"Generated: {timezone.now().strftime('%B %d, %Y')}"
        ])

        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Inches(1.1) # Date
        table.columns[1].width = Inches(4.1) # Description
        table.columns[2].width = Inches(1.3) # Amount
        table.columns[3].width = Inches(1.3) # Balance

        style_table_header(table.rows[0], ['Date', 'Description', 'Amount', 'Balance'])
        
        for row in final_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row['date'].strftime('%Y-%m-%d')
            row_cells[1].text = f"{row['desc']} ({row['ref']})"
            amt_str = f"{row['amount']:,.2f}"
            if row['is_credit']:
                amt_str = f"({amt_str})"
            row_cells[2].text = amt_str
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[3].text = f"{row['balance']:,.2f}"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"\nTotal Outstanding Balance: {running_balance:,.2f}").bold = True

        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    else: # PDF
        context = {
            'customer': customer,
            'ledger': final_data,
            'today': timezone.now(),
            'current_balance': running_balance
        }
        pdf = render_to_pdf('inventory/statement_pdf.html', context, request=request)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
            return response
        return HttpResponse("Error Generating PDF", status=500)

def generate_inventory_csv(products):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="inventory_snapshot.csv"'
    writer = csv.writer(response)
    writer.writerow(['Product', 'SKU', 'Category', 'Quantity', 'Price', 'Status'])
    for p in products:
        writer.writerow([p.name, p.sku, p.category.name if p.category else 'N/A', p.quantity, p.price, p.get_status_display()])
    return response

def generate_supplier_deliveries_export(supplier, purchase_orders, format_type, request):
    filename = f"Deliveries_{slugify(supplier.name)}_{timezone.now().strftime('%Y%m%d')}"

    if format_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        writer = csv.writer(response)
        writer.writerow(['PO ID', 'Order Date', 'Status', 'Product SKU', 'Product Name', 'Quantity', 'Price per Item'])
        for po in purchase_orders:
            for item in po.items.all():
                writer.writerow([po.order_id, po.order_date.strftime('%Y-%m-%d'), po.get_status_display(), item.product.sku, item.product.name, item.quantity, item.price])
        return response

    elif format_type == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.title = "Deliveries"
        ws.page_margins = PageMargins(left=0.25, right=0.25, top=0.25, bottom=0.25, header=0.3, footer=0.3)
        
        # Styles
        title_font = Font(name='Arial', size=18, bold=True, color="2C3E50")
        subtitle_font = Font(name='Arial', size=14, bold=True, color="2799A5")
        info_font = Font(name='Arial', size=10, color="7F8C8D")
        header_font = Font(name='Arial', size=10, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        po_header_fill = PatternFill(start_color="ECF0F1", end_color="ECF0F1", fill_type="solid")
        thin_border = Border(left=Side(style='thin', color="BDC3C7"), right=Side(style='thin', color="BDC3C7"), top=Side(style='thin', color="BDC3C7"), bottom=Side(style='thin', color="BDC3C7"))
        
        # Report Header
        add_excel_logo(ws)
        ws.merge_cells('B1:C1'); ws['B1'] = "Rich Land Auto Supply"; ws['B1'].font = title_font; ws['B1'].alignment = Alignment(horizontal='left', vertical='center')
        ws.merge_cells('D1:F1'); ws['D1'] = "SUPPLIER HISTORY"; ws['D1'].font = subtitle_font; ws['D1'].alignment = Alignment(horizontal='right', vertical='center')
        ws.merge_cells('D2:F2'); ws['D2'] = f"Supplier: {supplier.name}"; ws['D2'].font = info_font; ws['D2'].alignment = Alignment(horizontal='right')
        ws.merge_cells('D3:F3'); ws['D3'] = f"Generated: {timezone.now().strftime('%B %d, %Y %I:%M %p')}"; ws['D3'].font = info_font; ws['D3'].alignment = Alignment(horizontal='right')
        ws.row_dimensions[1].height = 30
        
        current_row = 5
        for po in purchase_orders:
            # PO Header
            ws.merge_cells(f'A{current_row}:C{current_row}'); cell = ws[f'A{current_row}']; cell.value = f"PO # {po.order_id}"; cell.font = Font(name='Arial', size=11, bold=True, color="2C3E50"); cell.alignment = Alignment(vertical='center', indent=1)
            ws.merge_cells(f'D{current_row}:E{current_row}'); cell = ws[f'D{current_row}']; cell.value = f"Date: {po.order_date.strftime('%Y-%m-%d')}"; cell.font = Font(name='Arial', size=10); cell.alignment = Alignment(vertical='center')
            cell = ws[f'F{current_row}']; cell.value = f"Status: {po.get_status_display()}"; cell.font = Font(name='Arial', size=10, bold=True); cell.alignment = Alignment(horizontal='right', vertical='center')
            for col in range(1, 7): ws.cell(row=current_row, column=col).border = thin_border; ws.cell(row=current_row, column=col).fill = po_header_fill
            current_row += 1
            
            # Table Headers
            headers = ['SKU', 'Product Name', 'Qty', 'Unit Price', 'Total', 'Rcv']
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=current_row, column=col_num); cell.value = header; cell.font = header_font; cell.fill = header_fill; cell.alignment = Alignment(horizontal='center', vertical='center'); cell.border = thin_border
            current_row += 1
            
            if not po.items.exists():
                ws.merge_cells(f'A{current_row}:F{current_row}'); cell = ws[f'A{current_row}']; cell.value = "No items found."; cell.alignment = Alignment(horizontal='center'); cell.border = thin_border
                current_row += 1
            else:
                for item in po.items.all():
                    ws.cell(row=current_row, column=1, value=item.product.sku).border = thin_border
                    ws.cell(row=current_row, column=2, value=item.product.name).border = thin_border
                    c3 = ws.cell(row=current_row, column=3, value=item.quantity); c3.alignment = Alignment(horizontal='center'); c3.border = thin_border
                    c4 = ws.cell(row=current_row, column=4, value=item.price); c4.number_format = '#,##0.00'; c4.border = thin_border
                    c5 = ws.cell(row=current_row, column=5, value=item.line_total); c5.number_format = '#,##0.00'; c5.font = Font(bold=True); c5.border = thin_border
                    c6 = ws.cell(row=current_row, column=6, value='Y' if po.status == 'RECEIVED' else '-'); c6.alignment = Alignment(horizontal='center'); c6.border = thin_border
                    current_row += 1
            current_row += 1
        
        ws.column_dimensions['A'].width = 15; ws.column_dimensions['B'].width = 45; ws.column_dimensions['C'].width = 10; ws.column_dimensions['D'].width = 12; ws.column_dimensions['E'].width = 15; ws.column_dimensions['F'].width = 8
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
        wb.save(response)
        return response

    elif format_type == 'word':
        document = Document()
        setup_word_document_margins(document)
        # Switch to Landscape
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11); section.page_height = Inches(8.5)

        create_header(document, "SUPPLIER HISTORY", [
            f"Supplier: {supplier.name}",
            f"Contact: {supplier.contact_person or 'N/A'}",
            f"Email: {supplier.email or 'N/A'}",
            f"Phone: {supplier.phone or 'N/A'}",
            f"Generated: {timezone.now().strftime('%B %d, %Y %I:%M %p')}"
        ], width_inches=5.5)

        for po in purchase_orders:
            # PO Header
            po_table = document.add_table(rows=1, cols=3); po_table.style = 'Table Grid'; po_table.autofit = False
            po_table.columns[0].width = Inches(2.5); po_table.columns[1].width = Inches(1.5); po_table.columns[2].width = Inches(1.5)
            row = po_table.rows[0]
            
            c1 = row.cells[0]; p = c1.paragraphs[0]; run = p.add_run(f"PO # {po.order_id}"); run.bold = True; run.font.name = 'Arial'; run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50); run.font.size = Pt(9); set_cell_background(c1, "ECF0F1")
            c2 = row.cells[1]; p = c2.paragraphs[0]; run = p.add_run(f"Date: {po.order_date.strftime('%Y-%m-%d')}"); run.font.name = 'Arial'; run.font.size = Pt(9); set_cell_background(c2, "ECF0F1")
            c3 = row.cells[2]; p = c3.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; run = p.add_run(f"Status: {po.get_status_display()}"); run.font.name = 'Arial'; run.bold = True; run.font.size = Pt(9); set_cell_background(c3, "ECF0F1")
            
            # Items
            table = document.add_table(rows=1, cols=6, style='Table Grid'); table.autofit = False
            widths = [Inches(0.8), Inches(2.0), Inches(0.5), Inches(0.8), Inches(1.0), Inches(0.4)]
            for i, width in enumerate(widths): table.columns[i].width = width
            
            hdr_cells = table.rows[0].cells
            headers = ['SKU', 'Product Name', 'Qty', 'Unit Price', 'Total', 'Rcv']
            alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER]
            for i, text in enumerate(headers):
                hdr_cells[i].text = ""; set_cell_background(hdr_cells[i], "2C3E50"); p = hdr_cells[i].paragraphs[0]; p.alignment = alignments[i]
                run = p.add_run(text); run.font.name = 'Arial'; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.bold = True; run.font.size = Pt(8)
            
            # Repeat header on new pages
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = parse_xml(r'<w:tblHeader %s/>' % nsdecls('w'))
            trPr.append(tblHeader)
            
            if not po.items.exists():
                row_cells = table.add_row().cells; row_cells[0].merge(row_cells[5]); row_cells[0].text = "No items found."; row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for item in po.items.all():
                    row_cells = table.add_row().cells
                    row_cells[0].text = item.product.sku; 
                    for run in row_cells[0].paragraphs[0].runs: run.font.name = 'Arial'; run.font.size = Pt(8)
                    row_cells[1].text = item.product.name; 
                    for run in row_cells[1].paragraphs[0].runs: run.font.name = 'Arial'; run.font.size = Pt(8)
                    row_cells[2].text = str(item.quantity); row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; 
                    for run in row_cells[2].paragraphs[0].runs: run.font.name = 'Arial'; run.font.size = Pt(8)
                    row_cells[3].text = f"{item.price:,.2f}"; row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; 
                    for run in row_cells[3].paragraphs[0].runs: run.font.name = 'Arial'; run.font.size = Pt(8)
                    row_cells[4].text = f"{item.line_total:,.2f}"; row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; 
                    for run in row_cells[4].paragraphs[0].runs: run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
                    row_cells[5].text = 'Y' if po.status == 'RECEIVED' else '-'; row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; 
                    for run in row_cells[5].paragraphs[0].runs: run.font.name = 'Arial'; run.font.size = Pt(8)
            document.add_paragraph()

        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    elif format_type == 'pdf':
        context = {'supplier': supplier, 'purchase_orders': purchase_orders, 'today': timezone.now(), 'request': request}
        pdf = render_to_pdf('inventory/supplier_deliveries_pdf.html', context)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
            return response
        return HttpResponse("Error Generating PDF", status=500)
    
    return None